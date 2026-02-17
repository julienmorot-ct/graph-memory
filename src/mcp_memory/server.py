# -*- coding: utf-8 -*-
"""
MCP Memory Server - Serveur principal.

Expose tous les outils MCP via HTTP/SSE avec FastMCP.
"""

import os
import sys
import json
import uuid
import base64
import argparse
from typing import Optional, List, Dict, Any

import uvicorn
from dotenv import load_dotenv

# Charger .env avant les imports qui en dépendent
load_dotenv()

from mcp.server.fastmcp import FastMCP, Context

from .config import get_settings
from .auth.middleware import AuthMiddleware, LoggingMiddleware, StaticFilesMiddleware, HostNormalizerMiddleware
from .auth.context import check_memory_access, check_write_permission, current_auth


# =============================================================================
# Initialisation
# =============================================================================

settings = get_settings()

# Créer l'instance FastMCP
# IMPORTANT: host="0.0.0.0" évite l'activation automatique de la protection
# DNS rebinding du SDK MCP v1.26+ qui n'autorise que localhost par défaut.
# Sans cela, les requêtes avec Host: graph-mem.mcp.cloud-temple.app sont
# rejetées avec un 421 Misdirected Request derrière un reverse proxy.
# Ref: mcp/server/fastmcp/server.py ligne 166 + mcp/server/transport_security.py
mcp = FastMCP(
    name=settings.mcp_server_name,
    host=settings.mcp_server_host,
    port=settings.mcp_server_port,
)


# =============================================================================
# Helpers - Services (lazy-loaded)
# =============================================================================

_graph_service = None
_storage_service = None
_extractor_service = None
_token_manager = None
_embedding_service = None
_chunker = None
_vector_store = None


def get_graph():
    """Lazy-load GraphService."""
    global _graph_service
    if _graph_service is None:
        from .core.graph import get_graph_service
        _graph_service = get_graph_service()
    return _graph_service


def get_storage():
    """Lazy-load StorageService."""
    global _storage_service
    if _storage_service is None:
        from .core.storage import get_storage_service
        _storage_service = get_storage_service()
    return _storage_service


def get_extractor():
    """Lazy-load ExtractorService."""
    global _extractor_service
    if _extractor_service is None:
        from .core.extractor import get_extractor_service
        _extractor_service = get_extractor_service()
    return _extractor_service


def get_tokens():
    """Lazy-load TokenManager."""
    global _token_manager
    if _token_manager is None:
        from .auth.token_manager import get_token_manager
        _token_manager = get_token_manager()
    return _token_manager


def get_embedder():
    """Lazy-load EmbeddingService."""
    global _embedding_service
    if _embedding_service is None:
        from .core.embedder import get_embedding_service
        _embedding_service = get_embedding_service()
    return _embedding_service


def get_chunker():
    """Lazy-load SemanticChunker."""
    global _chunker
    if _chunker is None:
        from .core.chunker import get_chunker as _get_chunker
        _chunker = _get_chunker()
    return _chunker


def get_vector_store():
    """Lazy-load VectorStoreService."""
    global _vector_store
    if _vector_store is None:
        from .core.vector_store import get_vector_store as _get_vs
        _vector_store = _get_vs()
    return _vector_store


_backup_service = None

def get_backup():
    """Lazy-load BackupService."""
    global _backup_service
    if _backup_service is None:
        from .core.backup import get_backup_service
        _backup_service = get_backup_service()
    return _backup_service


# =============================================================================
# OUTILS MCP - Gestion des Mémoires
# =============================================================================

@mcp.tool()
async def memory_create(
    memory_id: str,
    name: str,
    ontology: str,
    description: Optional[str] = None
) -> dict:
    """
    Crée une nouvelle mémoire (namespace isolé).
    
    L'ontologie est OBLIGATOIRE et copiée sur S3 pour persistance et versioning.
    
    Args:
        memory_id: Identifiant unique (ex: "quoteflow-legal")
        name: Nom lisible de la mémoire
        ontology: Nom de l'ontologie à utiliser (OBLIGATOIRE: legal, cloud, managed-services, technical)
        description: Description optionnelle
        
    Returns:
        Informations sur la mémoire créée
    """
    try:
        # Vérifier l'accès à la mémoire
        access_err = check_memory_access(memory_id)
        if access_err:
            return access_err
        
        # Vérifier que l'ontologie existe et la récupérer
        from .core.ontology import get_ontology_manager
        ontology_manager = get_ontology_manager()
        ontology_data = ontology_manager.get_ontology(ontology)
        
        if not ontology_data:
            available = [o["name"] for o in ontology_manager.list_ontologies()]
            return {
                "status": "error",
                "message": f"Ontologie '{ontology}' non trouvée. Disponibles: {available}"
            }
        
        # Stocker l'ontologie sur S3 pour la mémoire
        import yaml
        ontology_yaml = yaml.dump(ontology_data, allow_unicode=True, default_flow_style=False)
        ontology_bytes = ontology_yaml.encode('utf-8')
        
        ontology_s3_result = await get_storage().upload_document(
            memory_id=memory_id,
            filename=f"_ontology_{ontology}.yaml",
            content=ontology_bytes,
            metadata={"type": "ontology", "ontology_name": ontology}
        )
        
        print(f"📝 [Memory] Ontologie '{ontology}' stockée: {ontology_s3_result['uri']}", file=sys.stderr)
        
        # Créer la mémoire dans le graphe avec l'URI S3 de l'ontologie
        memory = await get_graph().create_memory(
            memory_id=memory_id,
            name=name,
            description=description,
            ontology=ontology,
            ontology_uri=ontology_s3_result["uri"]
        )
        
        return {
            "status": "created",
            "memory_id": memory.id,
            "name": memory.name,
            "description": memory.description,
            "ontology": memory.ontology,
            "ontology_uri": ontology_s3_result["uri"]
        }
    except ValueError as e:
        return {"status": "error", "message": str(e)}
    except Exception as e:
        return {"status": "error", "message": f"Erreur création: {str(e)}"}


@mcp.tool()
async def memory_delete(memory_id: str) -> dict:
    """
    Supprime une mémoire et tout son contenu (graphe + S3).
    
    ⚠️ ATTENTION: Cette opération est irréversible !
    Supprime le namespace Neo4j ET tous les fichiers S3 associés.
    
    Args:
        memory_id: ID de la mémoire à supprimer
        
    Returns:
        Statut de la suppression avec détails S3
    """
    try:
        # Vérifier l'accès à la mémoire
        access_err = check_memory_access(memory_id)
        if access_err:
            return access_err
        
        # 1. Supprimer la collection Qdrant (couplage strict)
        qdrant_deleted = False
        try:
            qdrant_deleted = await get_vector_store().delete_collection(memory_id)
        except Exception as e:
            print(f"❌ [Qdrant] Erreur suppression collection pour {memory_id}: {e}", file=sys.stderr)
            raise RuntimeError(f"Impossible de supprimer la collection Qdrant (couplage strict): {e}")
        
        # 2. Supprimer tous les fichiers S3 de la mémoire
        s3_result = {"deleted_count": 0, "error_count": 0}
        try:
            s3_result = await get_storage().delete_prefix(f"{memory_id}/")
            print(f"🗑️ [S3] Nettoyage mémoire {memory_id}: {s3_result['deleted_count']} fichiers supprimés", file=sys.stderr)
        except Exception as e:
            print(f"⚠️ [S3] Erreur nettoyage S3 pour {memory_id}: {e}", file=sys.stderr)
        
        # 3. Supprimer du graphe Neo4j
        deleted = await get_graph().delete_memory(memory_id)
        
        if deleted:
            return {
                "status": "deleted",
                "memory_id": memory_id,
                "qdrant_collection_deleted": qdrant_deleted,
                "s3_files_deleted": s3_result.get("deleted_count", 0),
                "s3_errors": s3_result.get("error_count", 0)
            }
        return {"status": "not_found", "memory_id": memory_id}
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool()
async def memory_list() -> dict:
    """
    Liste toutes les mémoires disponibles.
    
    Returns:
        Liste des mémoires avec leurs métadonnées
    """
    try:
        memories = await get_graph().list_memories()
        return {
            "status": "ok",
            "count": len(memories),
            "memories": [
                {
                    "id": m.id,
                    "name": m.name,
                    "description": m.description,
                    "ontology": m.ontology,
                    "created_at": m.created_at.isoformat() if m.created_at else None
                }
                for m in memories
            ]
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool()
async def memory_stats(memory_id: str) -> dict:
    """
    Récupère les statistiques d'une mémoire.
    
    Args:
        memory_id: ID de la mémoire
        
    Returns:
        Statistiques (documents, entités, relations, top entités)
    """
    try:
        # Vérifier l'accès à la mémoire
        access_err = check_memory_access(memory_id)
        if access_err:
            return access_err
        
        stats = await get_graph().get_memory_stats(memory_id)
        return {
            "status": "ok",
            "memory_id": memory_id,
            "document_count": stats.document_count,
            "entity_count": stats.entity_count,
            "relation_count": stats.relation_count,
            "top_entities": stats.top_entities
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}


# =============================================================================
# OUTILS MCP - Ingestion de Documents
# =============================================================================

@mcp.tool()
async def memory_ingest(
    memory_id: str,
    content_base64: str,
    filename: str,
    metadata: Optional[Dict[str, Any]] = None,
    force: bool = False,
    source_path: Optional[str] = None,
    source_modified_at: Optional[str] = None,
    ctx: Optional[Context] = None
) -> dict:
    """
    Ingère un document dans une mémoire.
    
    Le document est:
    1. Stocké sur S3
    2. Analysé par le LLM pour extraire entités/relations
    3. Les entités et relations sont ajoutées au graphe
    
    Métadonnées enrichies stockées sur le nœud Document :
    - hash SHA-256 (déduplication)
    - taille en bytes, longueur du texte extrait
    - type de fichier (extension)
    - chemin source et date de modification source (si fournis)
    - stats d'extraction (entités, relations, chunks)
    
    Args:
        memory_id: ID de la mémoire cible
        content_base64: Contenu du document encodé en base64
        filename: Nom du fichier
        metadata: Métadonnées additionnelles (optionnel)
        force: Si True, réingère même si le document existe déjà
        source_path: Chemin complet d'origine du fichier (ex: "legal/contracts/CGA.pdf")
        source_modified_at: Date de dernière modification du fichier source (ISO 8601, ex: "2026-01-15T10:30:00")
        
    Returns:
        Résultat de l'ingestion avec statistiques
    """
    try:
        import time as _time
        import gc
        _t0 = _time.monotonic()
        _steps_log = []
        
        def _mem_mb():
            """Retourne l'usage mémoire RSS du processus en MB."""
            try:
                import resource
                return resource.getrusage(resource.RUSAGE_SELF).ru_maxrss / (1024 * 1024)  # macOS = bytes
            except Exception:
                return 0
        
        # Helper pour logger les étapes (ctx.info si disponible + stderr)
        async def _log(msg):
            mem = _mem_mb()
            _steps_log.append({"t": round(_time.monotonic() - _t0, 1), "msg": msg})
            print(f"📋 [Ingest] {msg} [RSS={mem:.0f}MB]", file=sys.stderr)
            sys.stderr.flush()
            if ctx:
                try:
                    await ctx.info(msg)
                except Exception:
                    pass
        
        # Vérifier l'accès à la mémoire
        access_err = check_memory_access(memory_id)
        if access_err:
            return access_err
        
        # Décoder le contenu (libérer content_base64 ensuite — peut être volumineux)
        content = base64.b64decode(content_base64)
        content_size = len(content)
        await _log(f"📦 Décodage: {content_size} bytes ({filename})")
        del content_base64
        
        # Vérifier si la mémoire existe
        memory = await get_graph().get_memory(memory_id)
        if not memory:
            return {"status": "error", "message": f"Mémoire '{memory_id}' non trouvée"}
        
        # Calculer le hash pour déduplication
        doc_hash = get_storage().compute_hash(content)
        
        # Vérifier si déjà ingéré
        existing = await get_graph().get_document_by_hash(memory_id, doc_hash)
        if existing and not force:
            return {
                "status": "already_exists",
                "document_id": existing.id,
                "filename": existing.filename,
                "message": "Document déjà ingéré (utilisez force=true pour réingérer)"
            }
        
        # Si force=True et document existant, supprimer l'ancien d'abord
        if existing and force:
            await _log("🔄 Suppression de l'ancienne version...")
            delete_result = await get_graph().delete_document(memory_id, existing.id)
            print(f"🔄 [Ingest] Ancien supprimé: {delete_result.get('entities_deleted', 0)} entités orphelines, "
                  f"{delete_result.get('relations_deleted', 0)} relations", file=sys.stderr)
        
        # Upload vers S3
        await _log("📤 Upload S3...")
        s3_result = await get_storage().upload_document(
            memory_id=memory_id,
            filename=filename,
            content=content,
            metadata=metadata
        )
        await _log("✅ Upload S3 terminé")
        
        # Extraire le texte du document
        file_ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
        await _log(f"📄 Extraction texte ({file_ext})...")
        text = _extract_text(content, filename)
        
        if not text:
            return {
                "status": "warning",
                "message": "Document uploadé mais extraction texte impossible",
                "s3_uri": s3_result["uri"]
            }
        
        await _log(f"📄 Texte extrait: {len(text)} caractères")
        
        # Libérer les bytes bruts (on a le texte + déjà uploadé S3)
        del content
        gc.collect()
        
        # Extraction des entités/relations via LLM avec l'ontologie de la mémoire
        if not memory.ontology:
            return {
                "status": "error",
                "message": f"La mémoire '{memory_id}' n'a pas d'ontologie définie. "
                           f"Recréez-la avec une ontologie valide."
            }
        
        # Progress callback pour l'extracteur → route vers ctx.info()
        async def _extraction_progress(event: str, data: dict):
            if event == "extraction_start":
                mode = data.get("mode", "single")
                chunks_total = data.get("chunks_total", 1)
                text_len = data.get("text_length", 0)
                if mode == "chunked":
                    await _log(f"🔍 Extraction LLM: {chunks_total} chunks ({text_len} chars)")
                else:
                    await _log(f"🔍 Extraction LLM: 1 chunk ({text_len} chars)")
            elif event == "extraction_chunk_done":
                chunk = data.get("chunk", 0)
                total = data.get("chunks_total", 1)
                e_new = data.get("entities_new", 0)
                r_new = data.get("relations_new", 0)
                e_cum = data.get("entities_cumul", 0)
                r_cum = data.get("relations_cumul", 0)
                await _log(f"🔍 Chunk {chunk}/{total} terminé: +{e_new}E +{r_new}R (cumul: {e_cum}E {r_cum}R)")
        
        await _log(f"🔍 Démarrage extraction LLM (ontologie: {memory.ontology})...")
        extraction = await get_extractor().extract_with_ontology_chunked(
            text, memory.ontology, progress_callback=_extraction_progress
        )
        await _log(f"✅ Extraction terminée: {len(extraction.entities)} entités, {len(extraction.relations)} relations")
        
        # Déduire le type de fichier depuis l'extension
        file_ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
        
        # Créer le document dans le graphe avec métadonnées enrichies
        await _log("📊 Stockage dans le graphe Neo4j...")
        doc_id = str(uuid.uuid4())
        document = await get_graph().add_document(
            memory_id=memory_id,
            doc_id=doc_id,
            uri=s3_result["uri"],
            filename=filename,
            doc_hash=doc_hash,
            metadata=metadata,
            source_path=source_path,
            source_modified_at=source_modified_at,
            size_bytes=content_size,
            text_length=len(text),
            content_type=file_ext
        )
        
        # Ajouter les entités et relations
        graph_result = await get_graph().add_entities_and_relations(
            memory_id=memory_id,
            doc_id=doc_id,
            extraction=extraction
        )
        
        # === RAG Vectoriel : Chunking + Embedding + Qdrant (synchrone strict) ===
        await _log("🧩 Vectorisation RAG (chunking + embedding + Qdrant)...")
        chunks_stored = 0
        EMBED_BATCH_SIZE = 5  # Envoyer max 5 chunks par appel API embedding
        try:
            # S'assurer que la collection Qdrant existe
            await get_vector_store().ensure_collection(memory_id)
            await _log("🧩 Collection Qdrant prête")
            sys.stderr.flush()
            
            # Si force, supprimer les anciens chunks Qdrant
            if existing and force:
                await get_vector_store().delete_document_chunks(memory_id, existing.id)
                await _log("🧩 Anciens chunks supprimés")
                sys.stderr.flush()
            
            # Chunker le texte (CPU-bound → thread pool pour ne pas bloquer l'event loop)
            await _log("🧩 Chunking sémantique en cours...")
            sys.stderr.flush()
            import asyncio
            loop = asyncio.get_event_loop()
            chunks = await loop.run_in_executor(None, get_chunker().chunk_document, text, filename)
            await _log(f"🧩 Chunking terminé: {len(chunks)} chunks créés")
            sys.stderr.flush()
            
            if chunks:
                # Enrichir chaque chunk avec doc_id et memory_id
                for chunk in chunks:
                    chunk.doc_id = doc_id
                    chunk.memory_id = memory_id
                
                # Générer les embeddings par BATCHES (évite surcharge API)
                chunk_texts = [c.text for c in chunks]
                total_chunks = len(chunk_texts)
                all_embeddings = []
                
                for batch_start in range(0, total_chunks, EMBED_BATCH_SIZE):
                    batch_end = min(batch_start + EMBED_BATCH_SIZE, total_chunks)
                    batch_num = batch_start // EMBED_BATCH_SIZE + 1
                    total_batches = (total_chunks + EMBED_BATCH_SIZE - 1) // EMBED_BATCH_SIZE
                    batch_texts = chunk_texts[batch_start:batch_end]
                    
                    await _log(f"🔢 Embedding batch {batch_num}/{total_batches} ({len(batch_texts)} chunks)")
                    sys.stderr.flush()
                    
                    try:
                        batch_embeddings = await get_embedder().embed_texts(batch_texts)
                        all_embeddings.extend(batch_embeddings)
                        await _log(f"✅ Batch {batch_num}/{total_batches} OK ({len(all_embeddings)}/{total_chunks})")
                        sys.stderr.flush()
                    except Exception as embed_err:
                        print(f"❌ [Ingest] Erreur embedding batch {batch_num}: {embed_err}", file=sys.stderr)
                        sys.stderr.flush()
                        raise
                
                # Stocker dans Qdrant
                await _log(f"📦 Stockage Qdrant ({len(all_embeddings)} vecteurs)...")
                sys.stderr.flush()
                chunks_stored = await get_vector_store().store_chunks(
                    memory_id=memory_id,
                    doc_id=doc_id,
                    filename=filename,
                    chunks=chunks,
                    embeddings=all_embeddings
                )
                
                await _log(f"✅ RAG: {chunks_stored} chunks vectorisés")
                sys.stderr.flush()
        except Exception as e:
            # Couplage strict : si Qdrant échoue, on fait échouer l'ingestion
            print(f"❌ [Ingest] Erreur RAG vectoriel: {e}", file=sys.stderr)
            sys.stderr.flush()
            raise RuntimeError(f"Échec vectorisation Qdrant (couplage strict): {e}")
        
        # Compter les types de relations
        from collections import Counter
        relation_types = Counter(r.type for r in extraction.relations)
        entity_types = Counter(e.type for e in extraction.entities)
        
        _elapsed = round(_time.monotonic() - _t0, 1)
        await _log(f"🏁 Ingestion terminée en {_elapsed}s")
        
        return {
            "status": "ok",
            "document_id": doc_id,
            "filename": filename,
            "s3_uri": s3_result["uri"],
            "size_bytes": s3_result["size_bytes"],
            "entities_extracted": len(extraction.entities),
            "relations_extracted": len(extraction.relations),
            "entities_created": graph_result.get("entities_created", 0),
            "entities_merged": graph_result.get("entities_merged", 0),
            "relations_created": graph_result.get("relations_created", 0),
            "relations_merged": graph_result.get("relations_merged", 0),
            "entity_types": dict(entity_types),
            "relation_types": dict(relation_types),
            "chunks_stored": chunks_stored,
            "summary": extraction.summary,
            "key_topics": extraction.key_topics,
            "steps": _steps_log,
            "elapsed_seconds": _elapsed,
        }
        
    except Exception as e:
        print(f"❌ [Ingest] Erreur: {e}", file=sys.stderr)
        return {"status": "error", "message": str(e)}


def _extract_text(content: bytes, filename: str) -> Optional[str]:
    """
    Extrait le texte d'un document.
    
    Formats supportés: txt, md, html, docx, pdf, csv
    """
    ext = filename.lower().split('.')[-1] if '.' in filename else ''
    
    try:
        # Texte brut et Markdown
        if ext in ('txt', 'md'):
            return content.decode('utf-8', errors='ignore')
        
        # HTML
        elif ext in ('html', 'htm'):
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(content.decode('utf-8', errors='ignore'), 'html.parser')
            # Supprimer scripts et styles
            for script in soup(["script", "style"]):
                script.decompose()
            text = soup.get_text(separator='\n', strip=True)
            return text
        
        # PDF
        elif ext == 'pdf':
            from pypdf import PdfReader
            from io import BytesIO
            reader = PdfReader(BytesIO(content))
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            return "\n".join(text_parts)
        
        # DOCX (Word)
        elif ext == 'docx':
            from docx import Document
            from io import BytesIO
            doc = Document(BytesIO(content))
            
            text_parts = []
            
            # Extraire les paragraphes
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extraire le texte des tableaux
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            return "\n".join(text_parts)
        
        # CSV
        elif ext == 'csv':
            import csv
            from io import StringIO
            
            # Décoder le contenu
            text_content = content.decode('utf-8', errors='ignore')
            reader = csv.reader(StringIO(text_content))
            
            rows = []
            for row in reader:
                rows.append(" | ".join(row))
            
            return "\n".join(rows)
        
        else:
            # Tenter de décoder comme texte (fallback)
            return content.decode('utf-8', errors='ignore')
            
    except Exception as e:
        print(f"⚠️ [Extract] Erreur extraction texte ({ext}): {e}", file=sys.stderr)
        return None


# =============================================================================
# OUTILS MCP - Recherche
# =============================================================================

@mcp.tool()
async def memory_search(
    memory_id: str,
    query: str,
    limit: int = 10
) -> dict:
    """
    Recherche dans une mémoire (graph-first).
    
    Recherche les entités et documents correspondant à la requête.
    Utilise principalement le graphe, pas de RAG vectoriel.
    
    Args:
        memory_id: ID de la mémoire
        query: Requête de recherche
        limit: Nombre max de résultats
        
    Returns:
        Entités trouvées avec leurs documents liés
    """
    try:
        # Vérifier l'accès à la mémoire
        access_err = check_memory_access(memory_id)
        if access_err:
            return access_err
        
        # Recherche d'entités
        entities = await get_graph().search_entities(memory_id, search_query=query, limit=limit)
        
        # Pour chaque entité, récupérer le contexte complet
        results = []
        for entity in entities:
            context = await get_graph().get_entity_context(
                memory_id, entity["name"], depth=1
            )
            results.append({
                "entity": entity,
                "documents": context.documents,
                "related_entities": context.related_entities
            })
        
        return {
            "status": "ok",
            "query": query,
            "memory_id": memory_id,
            "result_count": len(results),
            "results": results
        }
        
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool()
async def question_answer(
    memory_id: str,
    question: str,
    limit: int = 10
) -> dict:
    """
    Pose une question sur une mémoire et obtient une réponse basée sur le graphe.
    
    Utilise le graphe de connaissances pour répondre à la question.
    Recherche les entités pertinentes puis génère une réponse avec le LLM.
    
    Args:
        memory_id: ID de la mémoire
        question: Question en langage naturel
        limit: Nombre max d'entités à rechercher (défaut: 10)
        
    Returns:
        Réponse générée avec les entités liées
    """
    try:
        # Vérifier l'accès à la mémoire
        access_err = check_memory_access(memory_id)
        if access_err:
            return access_err
        
        # 1. Rechercher les entités pertinentes dans le graphe
        print(f"🔎 [Q&A] Recherche graphe: memory={memory_id}, question='{question}', limit={limit}", file=sys.stderr)
        entities = await get_graph().search_entities(memory_id, search_query=question, limit=limit)
        
        if entities:
            entity_summary = ", ".join(f"{e['name']} ({e.get('type','?')})" for e in entities)
            print(f"📊 [Q&A] Graphe: {len(entities)} entités trouvées → {entity_summary}", file=sys.stderr)
        else:
            print(f"📊 [Q&A] Graphe: 0 entités trouvées → fallback RAG-only", file=sys.stderr)
        
        # 2. Récupérer le contexte de chaque entité + documents sources
        context_parts = []
        entity_names = []
        source_documents = {}  # doc_id -> {filename, id}

        for entity in entities:
            entity_names.append(entity["name"])
            ctx = await get_graph().get_entity_context(memory_id, entity["name"], depth=1)

            # Collecter les documents sources et les associer à l'entité
            entity_doc_names = []
            for doc in ctx.documents:
                if isinstance(doc, dict):
                    doc_id = doc.get('id', '')
                    doc_filename = doc.get('filename', doc_id)
                    if doc_id:
                        if doc_id not in source_documents:
                            source_documents[doc_id] = {
                                "id": doc_id,
                                "filename": doc_filename,
                            }
                        entity_doc_names.append(doc_filename)

            # Construire le contexte texte AVEC le document source
            doc_ref = f" [Source: {', '.join(entity_doc_names)}]" if entity_doc_names else ""
            ctx_text = f"- {entity['name']} ({entity.get('type', '?')}){doc_ref}"
            if entity.get('description'):
                ctx_text += f": {entity['description']}"

            for rel in ctx.relations:
                ctx_text += f"\n  → {rel.get('type', 'RELATED_TO')}: {rel.get('description', '')}"

            related = [r['name'] for r in ctx.related_entities]
            if related:
                ctx_text += f"\n  Lié à: {', '.join(related)}"

            context_parts.append(ctx_text)

        # 3. === RAG vectoriel : Graph-Guided si entités trouvées, sinon RAG-only ===
        rag_context_parts = []
        rag_chunks_used = 0
        rag_mode = "graph-guided" if entities else "rag-only"
        try:
            # Collecter les doc_ids identifiés par le graphe (vide si aucune entité)
            graph_doc_ids = list(source_documents.keys())

            # Vectoriser la question
            query_embedding = await get_embedder().embed_query(question)

            # Recherche Qdrant :
            # - Graph-Guided : filtrée par les documents identifiés par le graphe
            # - RAG-only : recherche sur TOUS les chunks de la mémoire (fallback)
            score_threshold = settings.rag_score_threshold
            chunk_limit = settings.rag_chunk_limit
            
            chunk_results = await get_vector_store().search(
                memory_id=memory_id,
                query_embedding=query_embedding,
                doc_ids=graph_doc_ids if graph_doc_ids else None,
                limit=chunk_limit
            )

            # Sauver tous les résultats avant filtrage (pour diagnostic)
            all_chunk_results = list(chunk_results)
            
            # Filtrer par seuil de score (en dessous = non pertinent)
            total_before = len(chunk_results)
            chunk_results = [cr for cr in chunk_results if cr.score >= score_threshold]
            filtered_out = total_before - len(chunk_results)

            # Construire le contexte RAG (chunks pertinents)
            for cr in chunk_results:
                rag_context_parts.append(cr.context_text)
                rag_chunks_used += 1
                # Ajouter les docs trouvés par RAG au source_documents
                if cr.chunk.doc_id and cr.chunk.doc_id not in source_documents:
                    source_documents[cr.chunk.doc_id] = {
                        "id": cr.chunk.doc_id,
                        "filename": cr.chunk.filename or "?",
                    }

            print(f"🔍 [Q&A] RAG ({rag_mode}): {rag_chunks_used} chunks retenus"
                  f" (seuil={score_threshold}, {filtered_out} filtrés sur {total_before})"
                  f"{f' | graph-guided: {len(graph_doc_ids)} docs' if graph_doc_ids else ' | tous documents'}", 
                  file=sys.stderr)
            
            # Log détaillé : score + section + aperçu texte de chaque chunk RETENU
            for i, cr in enumerate(chunk_results):
                section = cr.chunk.section_title or cr.chunk.article_number or "—"
                preview = cr.chunk.text[:80].replace('\n', ' ').strip()
                print(f"   📎 [{i+1}] score={cr.score:.4f} ✅ | {section} | \"{preview}...\"", file=sys.stderr)
            
            # Log des chunks FILTRÉS (sous le seuil) — diagnostic de pertinence RAG
            if filtered_out > 0:
                # Recalculer les chunks filtrés pour le log
                filtered_chunks = [cr for cr in all_chunk_results if cr.score < score_threshold]
                for i, cr in enumerate(filtered_chunks[:5]):  # Max 5 pour ne pas surcharger
                    section = cr.chunk.section_title or cr.chunk.article_number or "—"
                    preview = cr.chunk.text[:60].replace('\n', ' ').strip()
                    print(f"   📎 [F{i+1}] score={cr.score:.4f} ❌ | {section} | \"{preview}...\"", file=sys.stderr)

        except Exception as e:
            print(f"⚠️ [Q&A] Erreur RAG vectoriel: {e}", file=sys.stderr)
            # On continue avec le contexte graphe seul

        # Si ni le graphe ni le RAG n'ont trouvé quoi que ce soit → pas de contexte
        if not entities and rag_chunks_used == 0:
            return {
                "status": "ok",
                "answer": "Je n'ai pas trouvé d'informations pertinentes dans cette mémoire pour répondre à votre question.",
                "entities": [],
                "rag_chunks_used": 0,
                "source_documents": []
            }
        
        # 4. Construire la liste des documents pour le prompt
        doc_list = "\n".join(
            f"  - {doc['filename']}" for doc in source_documents.values()
        )
        
        # 5. Assembler le contexte final (graphe + RAG)
        graph_context = "\n".join(context_parts)
        rag_context = "\n\n".join(rag_context_parts) if rag_context_parts else ""
        
        # 6. Générer la réponse avec le LLM
        graph_ctx_len = len(graph_context) if graph_context else 0
        rag_ctx_len = len(rag_context) if rag_context else 0
        doc_count = len(source_documents)
        print(f"📝 [Q&A] Contexte LLM: graphe={graph_ctx_len} chars, RAG={rag_ctx_len} chars, docs={doc_count}", file=sys.stderr)
        
        prompt = f"""Tu es un assistant expert qui répond à des questions basées sur un graphe de connaissances et des extraits de documents.

Documents sources disponibles :
{doc_list}

=== CONTEXTE 1 : Graphe de connaissances (entités et relations) ===
{graph_context}

=== CONTEXTE 2 : Extraits de documents pertinents (RAG vectoriel) ===
{rag_context if rag_context else "(aucun extrait supplémentaire)"}

Question de l'utilisateur : {question}

CONSIGNES :
- Réponds de manière concise et précise en te basant UNIQUEMENT sur les contextes fournis.
- Privilégie les extraits de documents (CONTEXTE 2) pour les détails factuels et les citations.
- Utilise le graphe (CONTEXTE 1) pour la vue d'ensemble et les relations entre concepts.
- Cite systématiquement le document source quand tu affirmes quelque chose (ex: "Selon les CGA, …", "L'article X de la CGV prévoit que…").
- Si une information provient de plusieurs documents, précise lesquels.
- Si le contexte ne permet pas de répondre complètement, dis-le clairement.
- Utilise le format Markdown pour structurer ta réponse.
"""
        
        answer = await get_extractor().generate_answer(prompt)
        
        return {
            "status": "ok",
            "answer": answer,
            "entities": entity_names,
            "rag_chunks_used": rag_chunks_used,
            "source_documents": list(source_documents.values()),
            "context_used": graph_context
        }
        
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool()
async def memory_query(
    memory_id: str,
    query: str,
    limit: int = 10
) -> dict:
    """
    Interroge une mémoire et retourne les données structurées SANS génération LLM.
    
    Effectue la même recherche que question_answer (graphe + RAG vectoriel)
    mais retourne les données brutes structurées au lieu de crafter une réponse.
    Idéal pour les agents IA qui veulent construire leur propre réponse.
    
    Pipeline :
    1. Recherche d'entités dans le graphe (fulltext + CONTAINS)
    2. Récupération du contexte de chaque entité (voisins, relations, documents)
    3. Recherche RAG vectorielle (graph-guided ou rag-only)
    4. Retour des données structurées (pas d'appel LLM)
    
    Args:
        memory_id: ID de la mémoire
        query: Requête en langage naturel
        limit: Nombre max d'entités à rechercher (défaut: 10)
        
    Returns:
        Données structurées : entités, relations, chunks RAG, documents sources, stats
    """
    try:
        # Vérifier l'accès à la mémoire
        access_err = check_memory_access(memory_id)
        if access_err:
            return access_err
        
        # 1. Rechercher les entités pertinentes dans le graphe
        print(f"🔎 [Query] Recherche graphe: memory={memory_id}, query='{query}', limit={limit}", file=sys.stderr)
        entities = await get_graph().search_entities(memory_id, search_query=query, limit=limit)
        
        if entities:
            entity_summary = ", ".join(f"{e['name']} ({e.get('type','?')})" for e in entities)
            print(f"📊 [Query] Graphe: {len(entities)} entités trouvées → {entity_summary}", file=sys.stderr)
        else:
            print(f"📊 [Query] Graphe: 0 entités trouvées → fallback RAG-only", file=sys.stderr)
        
        # 2. Récupérer le contexte de chaque entité + documents sources
        enriched_entities = []
        source_documents = {}  # doc_id -> {filename, id}
        
        for entity in entities:
            ctx = await get_graph().get_entity_context(memory_id, entity["name"], depth=1)
            
            # Collecter les documents sources
            entity_docs = []
            for doc in ctx.documents:
                if isinstance(doc, dict):
                    doc_id = doc.get('id', '')
                    doc_filename = doc.get('filename', doc_id)
                    if doc_id:
                        if doc_id not in source_documents:
                            source_documents[doc_id] = {
                                "id": doc_id,
                                "filename": doc_filename,
                            }
                        entity_docs.append(doc_filename)
            
            # Construire l'entité enrichie
            enriched_entity = {
                "name": entity["name"],
                "type": entity.get("type", "?"),
                "description": entity.get("description", ""),
                "source_documents": entity_docs,
                "relations": [
                    {
                        "type": rel.get("type", "RELATED_TO"),
                        "target": rel.get("target", rel.get("to", "?")),
                        "description": rel.get("description", ""),
                    }
                    for rel in ctx.relations
                ],
                "related_entities": [
                    {
                        "name": r.get("name", r) if isinstance(r, dict) else str(r),
                        "type": r.get("type", "?") if isinstance(r, dict) else "?",
                    }
                    for r in ctx.related_entities
                ],
            }
            enriched_entities.append(enriched_entity)
        
        # 3. RAG vectoriel : Graph-Guided si entités, sinon RAG-only
        rag_chunks = []
        rag_mode = "graph-guided" if entities else "rag-only"
        rag_chunks_filtered = 0
        
        try:
            graph_doc_ids = list(source_documents.keys())
            query_embedding = await get_embedder().embed_query(query)
            
            score_threshold = settings.rag_score_threshold
            chunk_limit = settings.rag_chunk_limit
            
            chunk_results = await get_vector_store().search(
                memory_id=memory_id,
                query_embedding=query_embedding,
                doc_ids=graph_doc_ids if graph_doc_ids else None,
                limit=chunk_limit
            )
            
            total_before = len(chunk_results)
            retained = [cr for cr in chunk_results if cr.score >= score_threshold]
            rag_chunks_filtered = total_before - len(retained)
            
            for cr in retained:
                rag_chunks.append({
                    "text": cr.chunk.text,
                    "score": round(cr.score, 4),
                    "doc_id": cr.chunk.doc_id or "",
                    "filename": cr.chunk.filename or "?",
                    "section_title": cr.chunk.section_title or "",
                    "article_number": cr.chunk.article_number or "",
                    "chunk_index": cr.chunk.index if hasattr(cr.chunk, 'index') else 0,
                })
                # Ajouter les docs trouvés par RAG
                if cr.chunk.doc_id and cr.chunk.doc_id not in source_documents:
                    source_documents[cr.chunk.doc_id] = {
                        "id": cr.chunk.doc_id,
                        "filename": cr.chunk.filename or "?",
                    }
            
            print(f"🔍 [Query] RAG ({rag_mode}): {len(retained)} chunks retenus"
                  f" (seuil={score_threshold}, {rag_chunks_filtered} filtrés sur {total_before})", 
                  file=sys.stderr)
        
        except Exception as e:
            print(f"⚠️ [Query] Erreur RAG vectoriel: {e}", file=sys.stderr)
        
        # 4. Retourner les données structurées (PAS d'appel LLM)
        return {
            "status": "ok",
            "memory_id": memory_id,
            "query": query,
            "retrieval_mode": rag_mode,
            "entities": enriched_entities,
            "rag_chunks": rag_chunks,
            "source_documents": list(source_documents.values()),
            "stats": {
                "entities_found": len(enriched_entities),
                "rag_chunks_retained": len(rag_chunks),
                "rag_chunks_filtered": rag_chunks_filtered,
                "rag_score_threshold": settings.rag_score_threshold,
                "rag_chunk_limit": settings.rag_chunk_limit,
            },
        }
    
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool()
async def memory_get_context(
    memory_id: str,
    entity_name: str,
    depth: int = 1
) -> dict:
    """
    Récupère le contexte complet d'une entité.
    
    Retourne tout ce qu'on sait sur une entité:
    - Documents qui la mentionnent
    - Entités reliées
    - Types de relations
    
    Args:
        memory_id: ID de la mémoire
        entity_name: Nom de l'entité
        depth: Profondeur de traversée (1 = voisins directs)
        
    Returns:
        Contexte complet de l'entité
    """
    try:
        # Vérifier l'accès à la mémoire
        access_err = check_memory_access(memory_id)
        if access_err:
            return access_err
        
        context = await get_graph().get_entity_context(
            memory_id, entity_name, depth
        )
        
        return {
            "status": "ok",
            "entity_name": context.entity_name,
            "entity_type": context.entity_type,
            "depth": context.depth,
            "documents": context.documents,
            "related_entities": context.related_entities,
            "relations": context.relations
        }
        
    except Exception as e:
        return {"status": "error", "message": str(e)}


# =============================================================================
# OUTILS MCP - Admin / Tokens
# =============================================================================

@mcp.tool()
async def admin_create_token(
    client_name: str,
    permissions: Optional[List[str]] = None,
    memory_ids: Optional[List[str]] = None,
    expires_in_days: Optional[int] = None,
    email: Optional[str] = None
) -> dict:
    """
    Crée un nouveau token d'accès pour un client.
    
    ⚠️ Le token retourné ne sera affiché qu'une seule fois !
    
    Args:
        client_name: Nom du client (ex: "quoteflow")
        permissions: Permissions ["read", "write", "admin"]
        memory_ids: IDs des mémoires autorisées (vide = toutes)
        expires_in_days: Expiration en jours (optionnel)
        email: Adresse email du propriétaire (optionnel)
        
    Returns:
        Token généré (à conserver précieusement)
    """
    try:
        token = await get_tokens().create_token(
            client_name=client_name,
            permissions=permissions or ["read", "write"],
            memory_ids=memory_ids or [],
            expires_in_days=expires_in_days,
            email=email
        )
        
        return {
            "status": "ok",
            "client_name": client_name,
            "email": email,
            "token": token,
            "permissions": permissions or ["read", "write"],
            "memory_ids": memory_ids or [],
            "message": "⚠️ Conservez ce token, il ne sera plus affiché !"
        }
        
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool()
async def admin_list_tokens() -> dict:
    """
    Liste tous les tokens actifs.
    
    Note: Les tokens eux-mêmes ne sont pas affichés, seulement leurs métadonnées.
    
    Returns:
        Liste des tokens avec leurs infos
    """
    try:
        tokens = await get_tokens().list_tokens()
        
        return {
            "status": "ok",
            "count": len(tokens),
            "tokens": [
                {
                    "client_name": t.client_name,
                    "email": t.email,
                    "permissions": t.permissions,
                    "memory_ids": t.memory_ids,
                    "created_at": t.created_at.isoformat() if t.created_at else None,
                    "expires_at": t.expires_at.isoformat() if t.expires_at else None,
                    "token_hash": t.token_hash
                }
                for t in tokens
            ]
        }
        
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool()
async def admin_revoke_token(token_hash_prefix: str) -> dict:
    """
    Révoque un token.
    
    Args:
        token_hash_prefix: Début du hash du token (8+ caractères)
        
    Returns:
        Statut de la révocation
    """
    try:
        # Trouver le token par son préfixe
        tokens = await get_tokens().list_tokens(include_revoked=False)
        
        matching = [t for t in tokens if t.token_hash.startswith(token_hash_prefix)]
        
        if not matching:
            return {"status": "error", "message": "Token non trouvé"}
        
        if len(matching) > 1:
            return {"status": "error", "message": "Préfixe ambigu, soyez plus précis"}
        
        # Révoquer
        success = await get_tokens().revoke_token(matching[0].token_hash)
        
        if success:
            return {
                "status": "ok",
                "message": f"Token révoqué pour '{matching[0].client_name}'"
            }
        return {"status": "error", "message": "Échec révocation"}
        
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool()
async def admin_update_token(
    token_hash_prefix: str,
    add_memories: Optional[List[str]] = None,
    remove_memories: Optional[List[str]] = None,
    set_memories: Optional[List[str]] = None
) -> dict:
    """
    Met à jour les mémoires autorisées d'un token.
    
    Trois modes (mutuellement exclusifs avec set_memories) :
    - add_memories: Ajoute des mémoires à la liste existante
    - remove_memories: Retire des mémoires de la liste existante
    - set_memories: Remplace toute la liste ([] = accès à TOUTES les mémoires)
    
    Args:
        token_hash_prefix: Début du hash du token (8+ caractères)
        add_memories: Mémoires à ajouter (ex: ["JURIDIQUE", "CLOUD"])
        remove_memories: Mémoires à retirer (ex: ["JURIDIQUE"])
        set_memories: Remplacer toute la liste (ex: ["CLOUD"], ou [] pour tout autoriser)
        
    Returns:
        Anciennes et nouvelles mémoires autorisées
    """
    try:
        # Trouver le token par son préfixe
        tokens = await get_tokens().list_tokens(include_revoked=False)
        matching = [t for t in tokens if t.token_hash.startswith(token_hash_prefix)]
        
        if not matching:
            return {"status": "error", "message": "Token non trouvé"}
        
        if len(matching) > 1:
            return {"status": "error", "message": "Préfixe ambigu, soyez plus précis"}
        
        # Vérifier que les mémoires existent (si on en ajoute)
        memories_to_check = (add_memories or []) + (set_memories or [])
        if memories_to_check:
            existing_memories = await get_graph().list_memories()
            existing_ids = {m.id for m in existing_memories}
            unknown = [m for m in memories_to_check if m not in existing_ids]
            if unknown:
                return {
                    "status": "error",
                    "message": f"Mémoires inconnues: {unknown}. Disponibles: {sorted(existing_ids)}"
                }
        
        # Mettre à jour
        result = await get_tokens().update_token_memories(
            token_hash=matching[0].token_hash,
            add_memories=add_memories,
            remove_memories=remove_memories,
            set_memories=set_memories
        )
        
        if result:
            return {
                "status": "ok",
                "client_name": result["client_name"],
                "token_hash_prefix": result["token_hash"][:8] + "...",
                "previous_memories": result["previous_memories"],
                "current_memories": result["current_memories"],
                "message": (
                    "Accès à toutes les mémoires" if not result["current_memories"]
                    else f"Accès restreint à: {result['current_memories']}"
                )
            }
        return {"status": "error", "message": "Token non trouvé ou inactif"}
        
    except Exception as e:
        return {"status": "error", "message": str(e)}


# =============================================================================
# OUTILS MCP - Diagnostic
# =============================================================================

@mcp.tool()
async def memory_graph(memory_id: str, format: str = "full") -> dict:
    """
    Récupère le graphe complet d'une mémoire (entités, relations et documents).
    
    Utile pour visualiser ou exporter le graphe de connaissances.
    Inclut les documents avec leur URI S3 pour permettre la récupération.
    
    Args:
        memory_id: ID de la mémoire
        format: "full" (tout), "nodes" (entités+docs), "edges" (relations), "documents" (liste docs avec URI S3)
        
    Returns:
        nodes: Liste des entités et documents avec leurs propriétés
        edges: Liste des relations entre entités et documents
        documents: Liste des documents avec id, filename, uri S3
    """
    try:
        # Vérifier l'accès à la mémoire
        access_err = check_memory_access(memory_id)
        if access_err:
            return access_err
        
        graph_data = await get_graph().get_full_graph(memory_id)
        
        if format == "nodes":
            return {
                "status": "ok",
                "memory_id": memory_id,
                "node_count": len(graph_data["nodes"]),
                "nodes": graph_data["nodes"]
            }
        elif format == "edges":
            return {
                "status": "ok",
                "memory_id": memory_id,
                "edge_count": len(graph_data["edges"]),
                "edges": graph_data["edges"]
            }
        elif format == "documents":
            return {
                "status": "ok",
                "memory_id": memory_id,
                "document_count": len(graph_data["documents"]),
                "documents": graph_data["documents"]
            }
        else:  # full
            return {
                "status": "ok",
                "memory_id": memory_id,
                "node_count": len(graph_data["nodes"]),
                "edge_count": len(graph_data["edges"]),
                "document_count": len(graph_data["documents"]),
                "nodes": graph_data["nodes"],
                "edges": graph_data["edges"],
                "documents": graph_data["documents"]
            }
        
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool()
async def document_list(memory_id: str) -> dict:
    """
    Liste tous les documents d'une mémoire.
    
    Args:
        memory_id: ID de la mémoire
        
    Returns:
        Liste des documents avec leurs métadonnées
    """
    try:
        graph_data = await get_graph().get_full_graph(memory_id)
        docs = graph_data.get("documents", [])
        
        return {
            "status": "ok",
            "memory_id": memory_id,
            "count": len(docs),
            "documents": docs
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool()
async def document_get(
    memory_id: str,
    document_id: str,
    include_content: bool = False
) -> dict:
    """
    Récupère les métadonnées d'un document, et optionnellement son contenu.
    
    Par défaut, retourne uniquement les métadonnées (rapide, pas de téléchargement S3).
    Passez include_content=True pour télécharger et inclure le contenu du document.
    
    Args:
        memory_id: ID de la mémoire
        document_id: ID du document
        include_content: Si True, télécharge et inclut le contenu S3 (lent). Défaut: False.
        
    Returns:
        Métadonnées du document (et contenu si demandé)
    """
    try:
        # Récupérer les infos du document depuis le graphe (rapide, pas de S3)
        doc_info = await get_graph().get_document(memory_id, document_id)
        
        if not doc_info:
            return {"status": "error", "message": f"Document '{document_id}' non trouvé"}
        
        result = {
            "status": "ok",
            "document": {
                "id": doc_info.get("id"),
                "filename": doc_info.get("filename"),
                "uri": doc_info.get("uri"),
                "hash": doc_info.get("hash"),
                "ingested_at": doc_info.get("ingested_at"),
                "source_path": doc_info.get("source_path"),
                "source_modified_at": doc_info.get("source_modified_at"),
                "size_bytes": doc_info.get("size_bytes", 0),
                "text_length": doc_info.get("text_length", 0),
                "content_type": doc_info.get("content_type"),
            },
        }
        
        # Télécharger le contenu S3 seulement si demandé
        if include_content and doc_info.get("uri"):
            try:
                uri = doc_info["uri"]
                content_bytes = await get_storage().download_document(memory_id, uri)
                result["content"] = content_bytes.decode('utf-8', errors='ignore')
            except Exception as e:
                result["content"] = f"[Erreur lecture S3: {e}]"
        
        return result
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool()
async def document_delete(memory_id: str, document_id: str) -> dict:
    """
    Supprime un document du graphe ET de S3.
    
    Supprime :
    - Le fichier S3 associé
    - Le nœud Document dans Neo4j
    - Les relations MENTIONS du document
    - Les entités orphelines (non mentionnées par d'autres documents)
    - Les relations RELATED_TO impliquant des entités orphelines
    
    Args:
        memory_id: ID de la mémoire
        document_id: ID du document à supprimer
        
    Returns:
        Statut de la suppression avec compteurs (graphe + S3)
    """
    try:
        # Vérifier l'accès à la mémoire
        access_err = check_memory_access(memory_id)
        if access_err:
            return access_err
        
        # 1. Récupérer l'URI S3 avant suppression du graphe
        doc_info = await get_graph().get_document(memory_id, document_id)
        s3_deleted = False
        
        if doc_info and doc_info.get("uri"):
            # 2. Supprimer le fichier S3
            try:
                s3_deleted = await get_storage().delete_document(memory_id, doc_info["uri"])
                print(f"🗑️ [S3] Fichier supprimé: {doc_info['uri']}", file=sys.stderr)
            except Exception as e:
                print(f"⚠️ [S3] Erreur suppression S3 pour {doc_info['uri']}: {e}", file=sys.stderr)
        
        # 2b. Supprimer les chunks Qdrant (couplage strict)
        qdrant_chunks_deleted = 0
        try:
            qdrant_chunks_deleted = await get_vector_store().delete_document_chunks(memory_id, document_id)
        except Exception as e:
            print(f"❌ [Qdrant] Erreur suppression chunks pour doc {document_id}: {e}", file=sys.stderr)
            raise RuntimeError(f"Impossible de supprimer les chunks Qdrant (couplage strict): {e}")
        
        # 3. Supprimer du graphe Neo4j
        result = await get_graph().delete_document(memory_id, document_id)
        
        if result.get("deleted"):
            return {
                "status": "deleted",
                "document_id": document_id,
                "relations_deleted": result.get("relations_deleted", 0),
                "entities_deleted": result.get("entities_deleted", 0),
                "qdrant_chunks_deleted": qdrant_chunks_deleted,
                "s3_deleted": s3_deleted
            }
        return {"status": "error", "message": "Document non trouvé"}
        
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool()
async def ontology_list() -> dict:
    """
    Liste toutes les ontologies disponibles.
    
    Les ontologies définissent les règles d'extraction pour différents domaines.
    Chaque mémoire DOIT avoir une ontologie. Exemples:
    - legal: Documents juridiques et contractuels
    - cloud: Infrastructure cloud et certifications
    - managed-services: Infogérance et services managés
    - technical: Documentation technique et API
    
    Returns:
        Liste des ontologies avec leurs métadonnées
    """
    try:
        from .core.ontology import get_ontology_manager
        ontology_manager = get_ontology_manager()
        ontologies = ontology_manager.list_ontologies()
        
        return {
            "status": "ok",
            "count": len(ontologies),
            "ontologies": ontologies
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool()
async def storage_check(memory_id: Optional[str] = None) -> dict:
    """
    Vérifie la cohérence entre le graphe Neo4j et le stockage S3.
    
    Pour chaque mémoire (ou une mémoire spécifique) :
    1. Vérifie que chaque document du graphe est accessible sur S3
    2. Détecte les fichiers orphelins sur S3 (pas de référence dans le graphe)
    3. Retourne un rapport complet avec statistiques
    
    Args:
        memory_id: ID d'une mémoire spécifique (optionnel, toutes si omis)
        
    Returns:
        Rapport de cohérence S3/Graphe avec documents OK, manquants et orphelins
    """
    try:
        # 1. Récupérer les mémoires à vérifier
        if memory_id:
            memory = await get_graph().get_memory(memory_id)
            if not memory:
                return {"status": "error", "message": f"Mémoire '{memory_id}' non trouvée"}
            memories = [memory]
        else:
            memories = await get_graph().list_memories()
        
        # 2. Collecter toutes les URIs des documents référencés dans le graphe
        graph_uris = set()          # URIs référencées dans Neo4j
        graph_uri_details = {}      # URI -> {memory_id, filename, doc_id}
        memory_prefixes = set()     # Préfixes S3 des mémoires connues
        
        for mem in memories:
            mid = mem.id
            memory_prefixes.add(f"{mid}/")
            graph_data = await get_graph().get_full_graph(mid)
            
            for doc in graph_data.get("documents", []):
                uri = doc.get("uri", "")
                if uri:
                    graph_uris.add(uri)
                    graph_uri_details[uri] = {
                        "memory_id": mid,
                        "filename": doc.get("filename", "?"),
                        "doc_id": doc.get("id", "?")
                    }
        
        # 3. Vérifier l'accessibilité S3 de chaque document du graphe
        check_result = await get_storage().check_documents(list(graph_uris))
        
        # Enrichir les détails avec les infos du graphe
        for detail in check_result.get("details", []):
            uri = detail.get("uri", "")
            if uri in graph_uri_details:
                detail["memory_id"] = graph_uri_details[uri]["memory_id"]
                detail["filename"] = graph_uri_details[uri]["filename"]
                detail["doc_id"] = graph_uri_details[uri]["doc_id"]
        
        # 4. Lister tous les objets S3 pour détecter les orphelins
        #    IMPORTANT : pour la détection d'orphelins, on compare avec TOUTES
        #    les mémoires, pas seulement celles du scope. Sinon les docs des
        #    autres mémoires apparaissent comme faux-positifs.
        all_s3_objects = await get_storage().list_all_objects()
        
        # Collecter les clés S3 de TOUTES les mémoires (pas seulement le scope)
        all_graph_uris = set(graph_uris)  # Commencer avec celles du scope
        if memory_id:
            # Charger les URIs des autres mémoires aussi
            all_memories = await get_graph().list_memories()
            for mem in all_memories:
                if mem.id == memory_id:
                    continue  # Déjà chargé
                other_graph = await get_graph().get_full_graph(mem.id)
                for doc in other_graph.get("documents", []):
                    uri = doc.get("uri", "")
                    if uri:
                        all_graph_uris.add(uri)
        
        # Convertir les URIs du graphe en clés S3 pour comparaison
        graph_keys = set()
        for uri in all_graph_uris:
            try:
                key = get_storage()._parse_key(uri)
                graph_keys.add(key)
            except ValueError:
                pass
        
        # Ajouter les ontologies comme fichiers légitimes (pas orphelins)
        # Les fichiers _ontology_*.yaml sont des fichiers de config, pas des orphelins
        
        # Détecter les orphelins : sur S3 mais pas dans le graphe
        orphans = []
        for obj in all_s3_objects:
            key = obj["key"]
            
            # Ignorer les fichiers de health check
            if key.startswith("_health_check/"):
                continue
            
            # Ignorer les backups (gérés séparément via backup_list)
            if key.startswith("_backups/"):
                continue
            
            # Ignorer les ontologies (fichiers légitimes)
            # Le pattern est {hash[:8]}__ontology_{name}.yaml (double _ car hash + _ontology)
            if "_ontology_" in key:
                continue
            
            # Si la clé n'est pas référencée dans le graphe → orphelin
            if key not in graph_keys:
                orphans.append({
                    "key": key,
                    "uri": obj["uri"],
                    "size": obj["size"],
                    "last_modified": obj["last_modified"]
                })
        
        # 5. Construire le rapport
        def _human_size(size_bytes):
            """Convertit des bytes en taille lisible."""
            for unit in ['B', 'KB', 'MB', 'GB']:
                if size_bytes < 1024:
                    return f"{size_bytes:.1f} {unit}"
                size_bytes /= 1024
            return f"{size_bytes:.1f} TB"
        
        orphan_total_size = sum(o["size"] for o in orphans)
        
        report = {
            "status": "ok",
            "scope": memory_id or "all",
            "memories_checked": len(memories),
            "graph_documents": {
                "total": check_result["total"],
                "accessible": check_result["accessible"],
                "missing": check_result["missing"],
                "errors": check_result["errors"],
                "total_size": _human_size(check_result["total_size_bytes"]),
                "total_size_bytes": check_result["total_size_bytes"],
                "details": check_result["details"]
            },
            "s3_orphans": {
                "count": len(orphans),
                "total_size": _human_size(orphan_total_size),
                "total_size_bytes": orphan_total_size,
                "files": orphans
            },
            "s3_total_objects": len(all_s3_objects),
            "summary": (
                f"✅ {check_result['accessible']}/{check_result['total']} docs accessibles"
                + (f", ❌ {check_result['missing']} manquants" if check_result['missing'] > 0 else "")
                + (f", ⚠️ {len(orphans)} orphelins S3 ({_human_size(orphan_total_size)})" if orphans else "")
            )
        }
        
        return report
        
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool()
async def storage_cleanup(dry_run: bool = True) -> dict:
    """
    Nettoie les fichiers orphelins sur S3.
    
    Un fichier orphelin est un objet S3 qui n'est référencé par aucun document
    dans le graphe Neo4j (ni par une ontologie de mémoire).
    
    ⚠️ Par défaut, mode dry_run=True : liste les fichiers sans les supprimer.
    Passez dry_run=False pour effectuer la suppression.
    
    Args:
        dry_run: Si True, liste seulement. Si False, supprime réellement.
        
    Returns:
        Liste des fichiers orphelins (supprimés ou à supprimer)
    """
    try:
        # 1. Exécuter le check complet pour identifier les orphelins
        check = await storage_check()
        
        if check.get("status") != "ok":
            return check
        
        orphans = check.get("s3_orphans", {}).get("files", [])
        
        if not orphans:
            return {
                "status": "ok",
                "message": "Aucun fichier orphelin trouvé. Le S3 est propre ! 🧹",
                "orphans_found": 0,
                "deleted": 0,
                "dry_run": dry_run
            }
        
        if dry_run:
            return {
                "status": "ok",
                "message": f"🔍 {len(orphans)} fichiers orphelins trouvés ({check['s3_orphans']['total_size']}). "
                           f"Relancez avec dry_run=false pour les supprimer.",
                "orphans_found": len(orphans),
                "deleted": 0,
                "dry_run": True,
                "files": orphans
            }
        
        # 2. Supprimer les orphelins
        keys_to_delete = [o["key"] for o in orphans]
        delete_result = await get_storage().delete_objects(keys_to_delete)
        
        return {
            "status": "ok",
            "message": f"🗑️ {delete_result['deleted_count']} fichiers orphelins supprimés "
                       f"({check['s3_orphans']['total_size']} libérés).",
            "orphans_found": len(orphans),
            "deleted": delete_result["deleted_count"],
            "errors": delete_result["error_count"],
            "dry_run": False,
            "files": orphans
        }
        
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool()
async def system_health() -> dict:
    """
    Vérifie l'état de santé du système.
    
    Teste les connexions à tous les services :
    S3, Neo4j, LLMaaS, Qdrant, Embedding.
    Les 5 doivent être OK, sinon le service est considéré en erreur.
    
    Returns:
        État de chaque service
    """
    results = {}
    
    # Test S3
    try:
        results["s3"] = await get_storage().test_connection()
    except Exception as e:
        results["s3"] = {"status": "error", "message": str(e)}
    
    # Test Neo4j
    try:
        results["neo4j"] = await get_graph().test_connection()
    except Exception as e:
        results["neo4j"] = {"status": "error", "message": str(e)}
    
    # Test LLMaaS (génération)
    try:
        results["llmaas"] = await get_extractor().test_connection()
    except Exception as e:
        results["llmaas"] = {"status": "error", "message": str(e)}
    
    # Test Qdrant
    try:
        results["qdrant"] = await get_vector_store().test_connection()
    except Exception as e:
        results["qdrant"] = {"status": "error", "message": str(e)}
    
    # Test Embedding (LLMaaS endpoint)
    try:
        results["embedding"] = await get_embedder().test_connection()
    except Exception as e:
        results["embedding"] = {"status": "error", "message": str(e)}
    
    # Statut global : TOUS doivent être OK (couplage strict)
    all_ok = all(r.get("status") == "ok" for r in results.values())
    
    return {
        "status": "ok" if all_ok else "error",
        "services": results
    }


# =============================================================================
# OUTILS MCP - Backup / Restore
# =============================================================================

@mcp.tool()
async def backup_create(
    memory_id: str,
    description: Optional[str] = None,
    ctx: Optional[Context] = None
) -> dict:
    """
    Crée un backup complet d'une mémoire sur S3.
    
    Exporte le graphe Neo4j (entités, relations, documents),
    les vecteurs Qdrant (embeddings), et les références des documents S3.
    Applique la politique de rétention (BACKUP_RETENTION_COUNT).
    
    Args:
        memory_id: ID de la mémoire à sauvegarder
        description: Description optionnelle du backup
        
    Returns:
        backup_id, statistiques, temps d'exécution
    """
    try:
        # Sécurité : vérifier accès mémoire + permission write
        access_err = check_memory_access(memory_id)
        if access_err:
            return access_err
        write_err = check_write_permission()
        if write_err:
            return write_err
        
        async def _progress(msg):
            if ctx:
                try:
                    await ctx.info(msg)
                except Exception:
                    pass
        
        result = await get_backup().create_backup(
            memory_id=memory_id,
            description=description,
            progress_callback=_progress
        )
        return result
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool()
async def backup_list(memory_id: Optional[str] = None) -> dict:
    """
    Liste les backups disponibles sur S3.
    
    Args:
        memory_id: Si fourni, liste uniquement les backups de cette mémoire.
                   Sinon, liste tous les backups.
        
    Returns:
        Liste des backups avec date, taille, statistiques
    """
    try:
        backups = await get_backup().list_backups(memory_id=memory_id)
        
        return {
            "status": "ok",
            "count": len(backups),
            "backups": [
                {
                    "backup_id": b.get("backup_id"),
                    "memory_id": b.get("memory_id"),
                    "memory_name": b.get("memory_name"),
                    "description": b.get("description"),
                    "created_at": b.get("created_at"),
                    "stats": b.get("stats", {}),
                    "version": b.get("version"),
                }
                for b in backups
            ]
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool()
async def backup_restore(
    backup_id: str,
    ctx: Optional[Context] = None
) -> dict:
    """
    Restaure une mémoire depuis un backup S3.
    
    ⚠️ La mémoire NE DOIT PAS exister (erreur sinon).
    Supprimez-la d'abord avec memory_delete si nécessaire.
    
    Restaure le graphe Neo4j + les vecteurs Qdrant tels qu'ils étaient,
    SANS refaire l'extraction LLM (instantané).
    
    Args:
        backup_id: ID du backup (format: "memory_id/timestamp")
        
    Returns:
        Compteurs de restauration (entités, relations, vecteurs, documents S3)
    """
    try:
        # Sécurité : extraire memory_id du backup_id, vérifier accès + write
        from .core.backup import BackupService
        mid, _ = BackupService._validate_backup_id(backup_id)
        access_err = check_memory_access(mid)
        if access_err:
            return access_err
        write_err = check_write_permission()
        if write_err:
            return write_err
        
        async def _progress(msg):
            if ctx:
                try:
                    await ctx.info(msg)
                except Exception:
                    pass
        
        result = await get_backup().restore_backup(
            backup_id=backup_id,
            progress_callback=_progress
        )
        return result
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool()
async def backup_download(
    backup_id: str,
    include_documents: bool = False,
    ctx: Optional[Context] = None
) -> dict:
    """
    Télécharge un backup sous forme d'archive tar.gz encodée en base64.
    
    Par défaut (light) : uniquement les données JSON (graphe + vecteurs).
    Avec include_documents=True : inclut aussi les fichiers originaux (PDF, DOCX, etc.).
    
    Args:
        backup_id: ID du backup (format: "memory_id/timestamp")
        include_documents: Si True, inclut les documents originaux dans l'archive
        
    Returns:
        Archive tar.gz encodée en base64 + nom de fichier suggéré
    """
    try:
        # Sécurité : extraire memory_id du backup_id, vérifier accès mémoire
        from .core.backup import BackupService
        mid, _ = BackupService._validate_backup_id(backup_id)
        access_err = check_memory_access(mid)
        if access_err:
            return access_err
        
        async def _progress(msg):
            if ctx:
                try:
                    await ctx.info(msg)
                except Exception:
                    pass
        
        archive_bytes = await get_backup().download_backup(
            backup_id=backup_id,
            include_documents=include_documents,
            progress_callback=_progress
        )
        
        # Encoder en base64 pour transmission via MCP
        import base64 as b64
        archive_b64 = b64.b64encode(archive_bytes).decode("ascii")
        
        # Nom de fichier suggéré
        safe_id = backup_id.replace("/", "-")
        filename = f"backup-{safe_id}.tar.gz"
        
        return {
            "status": "ok",
            "backup_id": backup_id,
            "filename": filename,
            "size_bytes": len(archive_bytes),
            "include_documents": include_documents,
            "content_base64": archive_b64,
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool()
async def backup_delete(backup_id: str) -> dict:
    """
    Supprime un backup de S3.
    
    Args:
        backup_id: ID du backup (format: "memory_id/timestamp")
        
    Returns:
        Nombre de fichiers supprimés
    """
    try:
        # Sécurité : extraire memory_id du backup_id, vérifier accès + write
        from .core.backup import BackupService
        mid, _ = BackupService._validate_backup_id(backup_id)
        access_err = check_memory_access(mid)
        if access_err:
            return access_err
        write_err = check_write_permission()
        if write_err:
            return write_err
        
        result = await get_backup().delete_backup(backup_id)
        return result
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool()
async def backup_restore_archive(
    archive_base64: str,
    ctx: Optional[Context] = None
) -> dict:
    """
    Restaure une mémoire depuis une archive tar.gz (base64).
    
    L'archive doit contenir manifest.json, graph_data.json, qdrant_vectors.jsonl.
    Si elle contient un dossier documents/, les fichiers sont re-uploadés sur S3.
    
    ⚠️ La mémoire NE DOIT PAS exister (erreur sinon).
    
    Usage typique : backup download --include-documents → fichier.tar.gz → restore-file
    
    Args:
        archive_base64: Contenu de l'archive tar.gz encodé en base64
        
    Returns:
        Compteurs de restauration (entités, relations, vecteurs, documents S3)
    """
    try:
        # Sécurité : vérifier permission write avant restore
        write_err = check_write_permission()
        if write_err:
            return write_err
        
        archive_bytes = base64.b64decode(archive_base64)
        
        async def _progress(msg):
            if ctx:
                try:
                    await ctx.info(msg)
                except Exception:
                    pass
        
        result = await get_backup().restore_from_archive(
            archive_bytes=archive_bytes,
            progress_callback=_progress
        )
        return result
    except Exception as e:
        return {"status": "error", "message": str(e)}


# =============================================================================
# Point d'entrée
# =============================================================================

def main():
    """Point d'entrée principal."""
    parser = argparse.ArgumentParser(description="MCP Memory Server")
    parser.add_argument("--port", type=int, default=settings.mcp_server_port)
    parser.add_argument("--host", type=str, default=settings.mcp_server_host)
    parser.add_argument("--debug", action="store_true", default=settings.mcp_server_debug)
    args = parser.parse_args()
    
    # Récupérer l'app ASGI de FastMCP
    base_app = mcp.sse_app()
    
    # Empiler les middlewares (le dernier wrappé est le premier exécuté)
    # Flux requête : AuthMiddleware → LoggingMiddleware → StaticFilesMiddleware
    #                → HostNormalizerMiddleware → MCP SSE app
    #
    # HostNormalizerMiddleware : normalise le Host header pour que le MCP SDK
    # (Starlette) accepte les requêtes provenant de reverse proxies (nginx, Caddy)
    # qui transmettent le Host public (ex: "graph-mem.mcp.cloud-temple.app")
    # au lieu de "localhost:8002". Sans ce middleware, /sse et /messages
    # retournent HTTP 421 "Invalid Host header".
    app = HostNormalizerMiddleware(base_app)
    app = StaticFilesMiddleware(app)
    app = LoggingMiddleware(app, debug=args.debug)
    app = AuthMiddleware(app, debug=args.debug)
    
    # Afficher le banner
    print("=" * 70, file=sys.stderr)
    print("🧠 MCP Memory Server - Démarrage", file=sys.stderr)
    print(f"📡 Écoute sur http://{args.host}:{args.port}", file=sys.stderr)
    print(f"🔒 Auth     : Bearer Token (ou ADMIN_BOOTSTRAP_KEY)", file=sys.stderr)
    print(f"🐛 Debug    : {'ACTIVÉ' if args.debug else 'Désactivé'}", file=sys.stderr)
    print("=" * 70, file=sys.stderr)
    print("Outils disponibles:", file=sys.stderr)
    print("  - memory_create, memory_delete, memory_list, memory_stats", file=sys.stderr)
    print("  - memory_ingest, memory_search, memory_query, memory_get_context", file=sys.stderr)
    print("  - admin_create_token, admin_list_tokens, admin_revoke_token, admin_update_token", file=sys.stderr)
    print("  - storage_check, storage_cleanup, system_health", file=sys.stderr)
    print("  - backup_create, backup_list, backup_restore, backup_download, backup_delete", file=sys.stderr)
    print("=" * 70, file=sys.stderr)
    
    # Lancer le serveur
    uvicorn.run(app, host=args.host, port=args.port)


if __name__ == "__main__":
    main()
